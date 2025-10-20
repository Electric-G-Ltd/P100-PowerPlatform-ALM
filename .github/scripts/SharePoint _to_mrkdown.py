#!/usr/bin/env python3
"""
SharePoint Document to Markdown Converter

Converts .docx files from a SharePoint staging folder to markdown files
in the repository, with proper metadata preservation.

Usage:
    python scripts/sharepoint-to-markdown.py

Configuration:
    Update SOURCE_FOLDER and OUTPUT_FOLDER paths as needed.
"""

import os
import sys
import logging
from pathlib import Path
from datetime import datetime
from typing import Optional, List

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.shared import OxmlElement
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed.")
    print("Install it with: pip install python-docx")
    sys.exit(1)

# ============================================================================
# CONFIGURATION
# ============================================================================

# Source folder containing .docx files from SharePoint
SOURCE_FOLDER = r"C:\Dev P100 Word Staging folder from SharePoint"

# Destination folder for markdown files
OUTPUT_FOLDER = r"C:\Dev\P100-PowerPlatform-ALM\SharePoint"

# Enable debug logging
DEBUG = False

# ============================================================================
# LOGGING SETUP
# ============================================================================

LOG_LEVEL = logging.DEBUG if DEBUG else logging.INFO
logging.basicConfig(
    level=LOG_LEVEL,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler("sharepoint-to-markdown.log")
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def validate_paths() -> bool:
    """Validate that source and output folders exist or can be created."""
    source_path = Path(SOURCE_FOLDER)
    
    if not source_path.exists():
        logger.error(f"❌ Source folder not found: {SOURCE_FOLDER}")
        return False
    
    if not source_path.is_dir():
        logger.error(f"❌ Source path is not a directory: {SOURCE_FOLDER}")
        return False
    
    output_path = Path(OUTPUT_FOLDER)
    try:
        output_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"✅ Output folder ready: {OUTPUT_FOLDER}")
    except Exception as e:
        logger.error(f"❌ Cannot create output folder: {e}")
        return False
    
    return True


def get_docx_files() -> List[Path]:
    """Get all .docx files from source folder (non-recursive)."""
    source_path = Path(SOURCE_FOLDER)
    docx_files = list(source_path.glob("*.docx"))
    
    if not docx_files:
        logger.warning(f"⚠️  No .docx files found in {SOURCE_FOLDER}")
    else:
        logger.info(f"Found {len(docx_files)} .docx file(s)")
    
    return docx_files


def extract_document_metadata(doc) -> dict:
    """Extract metadata from a Word document."""
    metadata = {
        "title": None,
        "author": None,
        "created": None,
        "modified": None
    }
    
    # Get core properties
    try:
        props = doc.core_properties
        metadata["title"] = props.title or None
        metadata["author"] = props.author or None
        metadata["created"] = props.created.isoformat() if props.created else None
        metadata["modified"] = props.modified.isoformat() if props.modified else None
    except Exception as e:
        logger.warning(f"Could not extract metadata: {e}")
    
    return metadata


def get_document_title(doc, filename: str) -> str:
    """Extract or infer document title."""
    # Try to get title from document properties
    try:
        if doc.core_properties.title:
            return doc.core_properties.title
    except:
        pass
    
    # Try to get title from first paragraph (if it looks like a heading)
    try:
        if doc.paragraphs and len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0]
            if first_para.text.strip():
                return first_para.text.strip()
    except:
        pass
    
    # Fall back to filename (without extension)
    return Path(filename).stem


def paragraph_to_markdown(paragraph) -> Optional[str]:
    """Convert a Word paragraph to markdown text."""
    if not paragraph.text.strip():
        return None
    
    text = paragraph.text
    
    # Detect heading levels based on style name
    style = paragraph.style.name if hasattr(paragraph, 'style') else None
    
    if style and "Heading 1" in style:
        return f"# {text}"
    elif style and "Heading 2" in style:
        return f"## {text}"
    elif style and "Heading 3" in style:
        return f"### {text}"
    elif style and "Heading 4" in style:
        return f"#### {text}"
    elif style and "List" in style or style == "List Paragraph":
        return f"- {text}"
    else:
        return text


def docx_to_markdown(docx_path: Path) -> str:
    """Convert a .docx file to markdown content."""
    logger.info(f"Converting: {docx_path.name}")
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.error(f"❌ Failed to open {docx_path.name}: {e}")
        return ""
    
    # Extract metadata
    metadata = extract_document_metadata(doc)
    title = get_document_title(doc, docx_path.name)
    
    # Build markdown content
    md_lines = []
    
    # Add frontmatter
    md_lines.append("---")
    md_lines.append(f"title: {title}")
    if metadata["author"]:
        md_lines.append(f"author: {metadata['author']}")
    if metadata["created"]:
        md_lines.append(f"date_created: {metadata['created']}")
    if metadata["modified"]:
        md_lines.append(f"date_modified: {metadata['modified']}")
    md_lines.append(f"source_file: {docx_path.name}")
    md_lines.append(f"converted_at: {datetime.now().isoformat()}")
    md_lines.append("---")
    md_lines.append("")
    
    # Add document title as H1 if not already a heading
    if title and not title.startswith("#"):
        md_lines.append(f"# {title}")
        md_lines.append("")
    
    # Extract paragraphs
    for paragraph in doc.paragraphs:
        md_para = paragraph_to_markdown(paragraph)
        if md_para:
            md_lines.append(md_para)
    
    # Extract tables
    if doc.tables:
        md_lines.append("")
        md_lines.append("## Tables")
        md_lines.append("")
        for table in doc.tables:
            md_lines.extend(table_to_markdown(table))
            md_lines.append("")
    
    markdown_content = "\n".join(md_lines)
    return markdown_content


def table_to_markdown(table) -> List[str]:
    """Convert a Word table to markdown table format."""
    md_lines = []
    
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        md_lines.append("| " + " | ".join(cells) + " |")
        
        # Add separator after header row
        if row_idx == 0:
            separators = ["---"] * len(cells)
            md_lines.append("| " + " | ".join(separators) + " |")
    
    return md_lines


def save_markdown(content: str, output_path: Path) -> bool:
    """Save markdown content to file."""
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        logger.info(f"✅ Created: {output_path.name}")
        return True
    except Exception as e:
        logger.error(f"❌ Failed to save {output_path.name}: {e}")
        return False


# ============================================================================
# MAIN WORKFLOW
# ============================================================================

def main():
    """Main conversion workflow."""
    logger.info("=" * 70)
    logger.info("SharePoint Document to Markdown Converter")
    logger.info("=" * 70)
    logger.info(f"Source: {SOURCE_FOLDER}")
    logger.info(f"Output: {OUTPUT_FOLDER}")
    logger.info("")
    
    # Validate paths
    if not validate_paths():
        logger.error("❌ Path validation failed. Exiting.")
        return 1
    
    logger.info("")
    
    # Get .docx files
    docx_files = get_docx_files()
    if not docx_files:
        logger.warning("No documents to convert.")
        return 0
    
    logger.info("")
    
    # Convert each file
    successful = 0
    failed = 0
    
    for docx_path in sorted(docx_files):
        try:
            # Convert to markdown
            markdown_content = docx_to_markdown(docx_path)
            
            if not markdown_content:
                logger.warning(f"⚠️  Empty content from {docx_path.name}")
                failed += 1
                continue
            
            # Determine output filename
            output_filename = docx_path.stem + ".md"
            output_path = Path(OUTPUT_FOLDER) / output_filename
            
            # Save markdown file (will overwrite if exists)
            if save_markdown(markdown_content, output_path):
                successful += 1
            else:
                failed += 1
        
        except Exception as e:
            logger.error(f"❌ Unexpected error processing {docx_path.name}: {e}")
            failed += 1
    
    # Summary
    logger.info("")
    logger.info("=" * 70)
    logger.info(f"Conversion Complete: {successful} successful, {failed} failed")
    logger.info("=" * 70)
    
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())